import os
import markdown
import weasyprint
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MD_PATH = r"C:\project\smart\docs\tanitim.md"
PDF_PATH = r"C:\project\smart\docs\tanitim.pdf"
DOCX_PATH = r"C:\project\smart\docs\tanitim.docx"

with open(MD_PATH, encoding="utf-8") as f:
    md_text = f.read()

# ── HELPERS ─────────────────────────────────────────────────────────────────
CSS = """
@page { size: A4; margin: 2cm; }
body { font-family: 'Segoe UI', Arial, sans-serif; font-size: 11pt; line-height: 1.5; color: #1a1a1a; }
h1 { font-size: 22pt; color: #005a9e; border-bottom: 2px solid #005a9e; padding-bottom: 6px; margin-top: 30px; }
h2 { font-size: 16pt; color: #005a9e; margin-top: 28px; }
h3 { font-size: 13pt; color: #333; margin-top: 20px; }
h4 { font-size: 11pt; color: #444; margin-top: 14px; }
table { border-collapse: collapse; width: 100%; margin: 14px 0; page-break-inside: avoid; font-size: 9.5pt; }
th { background-color: #005a9e; color: white; padding: 7px 10px; text-align: left; font-weight: 700; border: 1px solid #005a9e; }
td { padding: 5px 10px; border: 1px solid #bbb; vertical-align: top; }
tr:nth-child(even) td { background-color: #f5f8fc; }
code { background: #eef; padding: 1px 5px; border-radius: 3px; font-family: 'Cascadia Code', 'Consolas', monospace; font-size: 9.5pt; }
pre { background: #f5f5f5; padding: 12px; border-left: 4px solid #005a9e; font-family: 'Cascadia Code', 'Consolas', monospace; font-size: 9pt; overflow-x: auto; }
hr { border: none; border-top: 1px solid #ddd; margin: 24px 0; }
strong { color: #005a9e; }
ul, ol { padding-left: 22px; }
li { margin: 3px 0; }
p { margin: 6px 0; }
"""


def set_cell_shading(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), color)
    tcPr.append(shading)


def add_styled_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = table._tbl.tblPr if table._tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "888888")
        tblBorders.append(el)
    tblPr.append(tblBorders)

    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_shading(cell, "005A9E")

    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            for p in cells[ci].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.5)
            if ri % 2 == 1:
                set_cell_shading(cells[ci], "F5F8FC")

    doc.add_paragraph()


def parse_table(lines):
    rows = []
    for raw in lines:
        s = raw.strip()
        if not s or s.startswith("|--"):
            continue
        cells = [c.strip() for c in s.strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return None
    headers = rows[0]
    data = rows[1:]
    return headers, data


# ── HTML → PDF ──────────────────────────────────────────────────────────────
html_body = markdown.markdown(md_text, extensions=["tables", "fenced_code"])
html_full = f"<!DOCTYPE html><html><head><meta charset='utf-8'><style>{CSS}</style></head><body>{html_body}</body></html>"

with open(r"C:\project\smart\docs\tanitim.html", "w", encoding="utf-8") as f:
    f.write(html_full)

weasyprint.HTML(string=html_full).write_pdf(PDF_PATH)
print(f"PDF created: {PDF_PATH}")

# ── MARKDOWN → DOCX ─────────────────────────────────────────────────────────
doc = Document()
style = doc.styles["Normal"]
style.font.name = "Segoe UI"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

lines = md_text.split("\n")
i = 0
while i < len(lines):
    raw = lines[i]
    s = raw.strip()
    i += 1

    if not s:
        continue

    if s.startswith("```"):
        code = []
        while i < len(lines) and not lines[i].strip().startswith("```"):
            code.append(lines[i])
            i += 1
        i += 1
        if code:
            p = doc.add_paragraph()
            run = p.add_run("\n".join(code))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        continue

    if s.startswith("# ") and not s.startswith("## "):
        m = re.match(r"^# (.+)", s)
        if m:
            doc.add_heading(m.group(1), 1)
        continue

    if s.startswith("## ") and not s.startswith("### "):
        m = re.match(r"^## (.+)", s)
        if m:
            doc.add_heading(m.group(1), 2)
        continue

    if s.startswith("### ") and not s.startswith("#### "):
        m = re.match(r"^### (.+)", s)
        if m:
            doc.add_heading(m.group(1), 3)
        continue

    if s.startswith("#### "):
        m = re.match(r"^#### (.+)", s)
        if m:
            doc.add_heading(m.group(1), 4)
        continue

    if s == "---":
        p = doc.add_paragraph()
        run = p.add_run("─" * 60)
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        continue

    # Table detection
    if "|" in s and s.startswith("|"):
        tbl_lines = [raw]
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt.startswith("|"):
                break
            tbl_lines.append(lines[i])
            i += 1
        parsed = parse_table(tbl_lines)
        if parsed:
            add_styled_table(doc, parsed[0], parsed[1])
        continue

    # Bullet list
    if s.startswith("- "):
        doc.add_paragraph(s[2:].lstrip(), style="List Bullet")
        continue

    # Regular paragraph — strip bold markdown markers
    clean = s.replace("**", "").replace("*", "")
    doc.add_paragraph(clean)

tmp_path = DOCX_PATH + ".tmp"
doc.save(tmp_path)
os.replace(tmp_path, DOCX_PATH)
print(f"DOCX created: {DOCX_PATH}")
